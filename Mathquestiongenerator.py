import random
from docx import Document
from docx.shared import Mm, Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def set_cell_style(cell):
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for border in ['top', 'left', 'bottom', 'right']:
        tag = OxmlElement(f'w:{border}')
        tag.set(qn('w:val'), 'single')
        tag.set(qn('w:sz'), '4')
        tag.set(qn('w:space'), '0')
        tag.set(qn('w:color'), 'auto')
        tcBorders.append(tag)
    tcPr.append(tcBorders)
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

def get_precedence(op):
    return 2 if op in ['*', '/'] else 1

def generate_qa():
    operators = ['+', '-', '*', '/']
    op1, op2 = random.choices(operators, k=2)
    
    while True:
        try:
            # Generate numbers based on operation precedence
            if get_precedence(op2) > get_precedence(op1):
                # Handle right operation first (b op2 c)
                if op2 == '+':
                    b = random.randint(1, 50)
                    c = random.randint(1, 50)
                    intermediate = b + c
                elif op2 == '-':
                    b = random.randint(2, 100)
                    c = random.randint(1, b-1)
                    intermediate = b - c
                elif op2 == '*':
                    b = random.randint(1, 12)
                    c = random.randint(1, 12)
                    intermediate = b * c
                elif op2 == '/':
                    c = random.randint(1, 10)
                    b = c * random.randint(1, 10)
                    intermediate = b // c

                # Generate a for first operation
                if op1 == '+':
                    a = random.randint(1, 100)
                    answer = a + intermediate
                elif op1 == '-':
                    a = random.randint(intermediate, intermediate + 100)
                    answer = a - intermediate
                elif op1 == '*':
                    a = random.randint(1, 10)
                    answer = a * intermediate
                elif op1 == '/':
                    a = intermediate * random.randint(1, 10)
                    answer = a // intermediate
            else:
                # Handle left operation first (a op1 b)
                if op1 == '+':
                    a = random.randint(1, 100)
                    b = random.randint(1, 100)
                    intermediate = a + b
                elif op1 == '-':
                    a = random.randint(2, 100)
                    b = random.randint(1, a-1)
                    intermediate = a - b
                elif op1 == '*':
                    a = random.randint(1, 12)
                    b = random.randint(1, 12)
                    intermediate = a * b
                elif op1 == '/':
                    b = random.randint(1, 10)
                    a = b * random.randint(1, 10)
                    intermediate = a // b

                # Generate c for second operation
                if op2 == '+':
                    c = random.randint(1, 100)
                    answer = intermediate + c
                elif op2 == '-':
                    c = random.randint(1, intermediate)
                    answer = intermediate - c
                elif op2 == '*':
                    c = random.randint(1, 10)
                    answer = intermediate * c
                elif op2 == '/':
                    c = random.choice([n for n in range(1, intermediate+1) if intermediate % n == 0])
                    answer = intermediate // c

            # Validate final answer
            if answer > 0 and isinstance(answer, int):
                break
        except:
            continue

    question = f"{a} {op1} {b} {op2} {c} = ___"
    return question, answer

def create_worksheet():
    doc = Document()
    
    # Page setup
    section = doc.sections[0]
    section.page_height = Mm(297)
    section.page_width = Mm(210)
    section.left_margin = Mm(20)
    section.right_margin = Mm(20)
    
    # Set default font
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(14)
    
    # Create 10 pages
    for page in range(10):
        # Generate questions and answers
        qa_pairs = [generate_qa() for _ in range(20)]
        
        # Create questions table
        q_table = doc.add_table(rows=10, cols=2)
        q_table.style = 'Table Grid'
        
        for i in range(10):
            for col in range(2):
                cell = q_table.cell(i, col)
                index = i + col*10
                cell.text = f"{index+1}. {qa_pairs[index][0]}"
                set_cell_style(cell)
        
        # Add space between sections
        doc.add_paragraph("\n" * 8)
        
        # Create answers table
        a_table = doc.add_table(rows=10, cols=2)
        a_table.style = 'Table Grid'
        for i in range(10):
            for col in range(2):
                cell = a_table.cell(i, col)
                index = i + col*10
                cell.text = f"{index+1}. {qa_pairs[index][1]}"
                set_cell_style(cell)
                # Add answer styling
                shading = OxmlElement('w:shd')
                shading.set(qn('w:fill'), 'F0F0F0')
                cell._tc.get_or_add_tcPr().append(shading)
        
        if page < 9:
            doc.add_page_break()
    
    doc.save('math_worksheet.docx')

if __name__ == "__main__":
    create_worksheet()